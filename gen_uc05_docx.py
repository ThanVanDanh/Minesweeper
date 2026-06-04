# -*- coding: utf-8 -*-
"""
Generate UC05 specification document as a professional .docx file.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

# ── Styles helpers ────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    sizes = {1: 14, 2: 13, 3: 12}
    set_font(run, bold=True, size=sizes.get(level, 12))
    return p

def add_body(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, bold=bold, size=11)
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "4472C4")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def styled_cell(cell, text, bold=False, bg=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                color=None, size=11, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        set_cell_bg(cell, bg)
    set_cell_borders(cell)
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size, color=color, italic=italic)

def make_table(doc, rows, cols, col_widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    return table

# ══════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IV. MÔ TẢ USE CASE UC05")
set_font(run, bold=True, size=14)

# ══════════════════════════════════════════════════════════════
# SECTION 1: USE CASE OVERVIEW TABLE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
table = make_table(doc, 10, 2, col_widths=[4.5, 12.5])

HEADER_BG  = "1F3864"   # dark navy
ROW1_BG    = "D6E4F7"   # light blue
ROW2_BG    = "FFFFFF"   # white

# Row 0: Table header
styled_cell(table.cell(0,0), "Thuộc tính",
            bold=True, bg=HEADER_BG, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
styled_cell(table.cell(0,1), "Mô tả",
            bold=True, bg=HEADER_BG, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)

overview_rows = [
    ("Usecase ID",    "UC05"),
    ("Tên Usecase",   "Quản lý Người chơi và Thành tích"),
    ("Tác nhân",      "Admin"),
    ("Trigger",       "Admin nhấn chọn mục \"Quản lý người chơi\" hoặc \"Quản lý kết quả\" từ thanh điều hướng."),
    ("Mô tả",
     "Cho phép Admin theo dõi toàn bộ danh sách người dùng; tìm kiếm nâng cao theo từ khoá "
     "(Username / Tên hiển thị), lọc theo email, vai trò, trạng thái, ngày tạo, lần đăng nhập "
     "gần nhất và số ván đã chơi; đồng thời thực hiện các thao tác: Thêm mới, Chỉnh sửa (Nickname, "
     "Vai trò), Khoá / Mở khoá và Xoá tài khoản. Ngoài ra, Admin có thể quản lý kết quả game, tự "
     "động phát hiện gian lận, xem thống kê tổng hợp theo người chơi và gửi thông báo đến người dùng."),
    ("Pre-condition", "Admin đã đăng nhập thành công vào hệ thống với vai trò ADMIN."),
    ("Post-condition",
     "- Các thay đổi về thông tin hoặc trạng thái người dùng được cập nhật chính xác vào CSDL.\n"
     "- Các kết quả được chọn bị xóa hoàn toàn khỏi CSDL.\n"
     "- Thông báo gửi đến người dùng được lưu trong bảng notifications.\n"
     "- Mọi thao tác quản trị được ghi nhận đầy đủ trong audit_log."),
    ("Phiên bản",     "v1.2 — 04/06/2026"),
]

for i, (label, value) in enumerate(overview_rows):
    bg = ROW1_BG if i % 2 == 0 else ROW2_BG
    styled_cell(table.cell(i+1, 0), label, bold=True, bg=bg)
    styled_cell(table.cell(i+1, 1), value, bg=bg)

# Caption
p = doc.add_paragraph("Bảng 10: Mô tả Use Case UC05 (phiên bản cập nhật v1.2)")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]; set_font(run, italic=True, size=10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# SECTION 2: BASIC FLOW
# ══════════════════════════════════════════════════════════════
add_heading(doc, "Basic Flow — UC05.1: Xem danh sách người chơi", 2)

basic_steps = [
    ("05.1.1", "Admin nhấn chọn mục \"Quản lý người chơi\" từ thanh điều hướng."),
    ("05.1.2", "Hệ thống hiển thị màn hình quản lý người chơi và các nút chức năng "
               "(Thêm, Chỉnh sửa, Khoá/Mở khoá, Xoá, Làm mới, Xem nhật ký)."),
    ("05.1.3", "Hệ thống tự động tải danh sách người chơi; hiển thị tối đa 20 người mỗi trang "
               "cùng thanh điều hướng phân trang (trang hiện tại / tổng số trang)."),
    ("05.1.4", "[CẢI TIẾN v1.2 – D2] Cột \"Trạng thái\" hiển thị badge màu trực quan: "
               "🟢 Hoạt động / 🔴 Đã khoá thay vì chỉ hiển thị văn bản thuần."),
    ("05.1.5", "[CẢI TIẾN v1.2 – D2] Panel thống kê tổng quan hiển thị: Tổng người dùng | "
               "Đang hoạt động | Đã khoá | Số Admin."),
    ("05.1.6", "Admin chọn một trong các chức năng: Tìm kiếm & Lọc nâng cao, Thêm, Chỉnh sửa, "
               "Khoá/Mở khoá, Xoá hoặc Xem lịch sử thao tác."),
]

table_bf = make_table(doc, len(basic_steps)+1, 2, col_widths=[3.0, 14.0])
styled_cell(table_bf.cell(0,0), "Bước",   bold=True, bg=HEADER_BG, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
styled_cell(table_bf.cell(0,1), "Nội dung", bold=True, bg=HEADER_BG, color=(255,255,255))
for i, (step, desc) in enumerate(basic_steps):
    bg = ROW1_BG if i % 2 == 0 else ROW2_BG
    styled_cell(table_bf.cell(i+1,0), step, bold=True, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_cell(table_bf.cell(i+1,1), desc, bg=bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# HELPER: Build a flow table
# ══════════════════════════════════════════════════════════════
def add_flow_table(doc, title, rows, badge=None):
    """rows = list of (step_id, description) or section headers (None, bold_text)."""
    full_title = title
    if badge:
        full_title = f"{title}  [{badge}]"
    add_heading(doc, full_title, 2)

    n = len(rows)
    t = make_table(doc, n+1, 2, col_widths=[3.0, 14.0])
    styled_cell(t.cell(0,0), "Bước",   bold=True, bg=HEADER_BG, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_cell(t.cell(0,1), "Nội dung", bold=True, bg=HEADER_BG, color=(255,255,255))

    SECTION_BG = "2E75B6"
    for i, (sid, desc) in enumerate(rows):
        if sid is None:  # section header
            styled_cell(t.cell(i+1,0), "",     bg=SECTION_BG)
            styled_cell(t.cell(i+1,1), desc, bold=True, bg=SECTION_BG, color=(255,255,255))
        else:
            bg = ROW1_BG if i % 2 == 0 else ROW2_BG
            styled_cell(t.cell(i+1,0), sid,  bold=True, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
            styled_cell(t.cell(i+1,1), desc, bg=bg)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# UC05.2 – Tìm kiếm nâng cao
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.2: Tìm kiếm & Lọc nâng cao người dùng",
               badge="CẢI TIẾN v1.1",
rows=[
    (None, "Luồng chính"),
    ("05.2.1", "Admin nhập từ khoá vào ô tìm kiếm (Username / Tên hiển thị) và/hoặc điền các "
               "bộ lọc mở rộng: Email, Vai trò, Trạng thái, Khoảng ngày tạo, Khoảng lần đăng "
               "nhập gần nhất, Số ván tối thiểu / tối đa, Có/không có kết quả game."),
    ("05.2.2", "Admin nhấn nút Tìm kiếm. Hệ thống kiểm tra hợp lệ (ngày From ≤ To, số ván ≥ 0, "
               "min ≤ max). Nếu sai, hiển thị thông báo lỗi cụ thể và dừng."),
    ("05.2.3", "Hệ thống gửi toàn bộ điều kiện xuống CSDL, lấy danh sách phù hợp."),
    ("05.2.4", "Hệ thống hiển thị từ trang đầu tiên và cập nhật nhãn \"Tìm thấy N kết quả\"."),
    (None, "Luồng thay thế"),
    ("05.2-A1", "Admin xoá hết điều kiện rồi nhấn Tìm kiếm → Hệ thống tải lại toàn bộ danh sách "
                "không kèm bộ lọc, hiển thị từ trang đầu tiên."),
    ("05.2-A2", "Admin nhấn Làm mới → Hệ thống xoá ô tìm kiếm, reset tất cả bộ lọc về mặc định "
                "\"Tất cả\" / trống, truy vấn lại và hiển thị toàn bộ danh sách từ trang đầu."),
])

# ══════════════════════════════════════════════════════════════
# UC05.3 – Thêm người dùng
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.3: Thêm người dùng",
               badge="CẢI TIẾN v1.2 – D1",
rows=[
    (None, "Luồng chính"),
    ("05.3.1", "Admin nhấn nút \"Thêm người chơi\"."),
    ("05.3.2", "Hệ thống hiển thị dialog nhập thông tin: Username (*), Nickname, Mật khẩu (mặc định 123456), "
               "Vai trò."),
    ("05.3.3", "Admin nhập thông tin và nhấn nút Thêm."),
    ("05.3.4", "Hệ thống hash mật khẩu (MD5) và lưu người dùng vào CSDL."),
    ("05.3.5", "Hệ thống gán ID trả về cho User mới, chuyển đến trang cuối cùng và làm nổi bật dòng mới."),
    ("05.3.6", "Hệ thống ghi nhận Log vào CSDL (action: CREATE_USER)."),
    (None, "Luồng thay thế"),
    ("05.3-A1", "Admin nhấn Huỷ → [CẢI TIẾN v1.2 – D1] Nếu Admin đã nhập thông tin, hệ thống hiển thị "
                "cảnh báo \"Bạn có muốn bỏ thay đổi không?\" trước khi đóng dialog."),
    ("05.3-A2", "Nickname bị để trống → Hệ thống tự động lấy Username làm tên hiển thị."),
])

# ══════════════════════════════════════════════════════════════
# UC05.4 – Chỉnh sửa
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.4: Chỉnh sửa thông tin người dùng",
               badge="CẢI TIẾN v1.2 – D1",
rows=[
    (None, "Luồng chính"),
    ("05.4.1", "Admin chọn user trong bảng."),
    ("05.4.2", "Admin nhấn nút Chỉnh sửa."),
    ("05.4.3", "Hệ thống hiển thị dialog sửa người dùng với thông tin hiện tại được điền sẵn "
               "(Username ở chế độ read-only, password ẩn)."),
    ("05.4.4", "Admin cập nhật Nickname và/hoặc Vai trò rồi nhấn Cập nhật."),
    ("05.4.5", "Hệ thống cập nhật thông tin vào CSDL."),
    ("05.4.6", "Hệ thống load lại danh sách, giữ nguyên trang hiện tại."),
    ("05.4.7", "Hệ thống ghi nhận Log vào CSDL với nội dung thay đổi trước/sau (action: UPDATE_USER)."),
    (None, "Luồng thay thế"),
    ("05.4-A1", "Admin nhấn Huỷ → [CẢI TIẾN v1.2 – D1] Nếu Admin đã thay đổi thông tin, hệ thống "
                "hiển thị cảnh báo \"Bạn có muốn bỏ thay đổi không?\" trước khi đóng dialog."),
])

# ══════════════════════════════════════════════════════════════
# UC05.5 – Khoá / Mở khoá
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.5: Khoá / Mở khoá tài khoản",
               badge="CẢI TIẾN v1.1 – A3, A4",
rows=[
    (None, "Luồng chính"),
    ("05.5.1", "Admin chọn User trong bảng và nhấn nút Khoá / Mở khoá."),
    ("05.5.2", "[CẢI TIẾN v1.1 – A4] Hệ thống kiểm tra: nếu User được chọn chính là Admin đang "
               "đăng nhập, hiển thị lỗi \"Bạn không thể khoá tài khoản của chính mình!\" và dừng."),
    ("05.5.3", "[CẢI TIẾN v1.1 – A3] Hệ thống hiển thị dialog xác nhận:\n"
               "  • Khoá: \"Tài khoản bị khoá sẽ không thể đăng nhập.\"\n"
               "  • Mở khoá: \"Tài khoản sẽ được phép đăng nhập trở lại.\""),
    ("05.5.4", "Admin xác nhận OK → Hệ thống cập nhật trạng thái is_active vào CSDL."),
    ("05.5.5", "Hệ thống load lại danh sách, cột Trạng thái cập nhật badge màu tương ứng."),
    ("05.5.6", "Hệ thống ghi nhận Log vào CSDL (action: LOCK_USER / UNLOCK_USER)."),
    (None, "Luồng thay thế"),
    ("05.5-A1", "Admin nhấn Huỷ tại dialog xác nhận → Đóng dialog, không thay đổi trạng thái."),
])

# ══════════════════════════════════════════════════════════════
# UC05.6 – Xóa người dùng
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.6: Xóa người dùng",
               badge="CẢI TIẾN v1.1 – A4",
rows=[
    (None, "Luồng chính"),
    ("05.6.1", "Admin chọn User trong bảng và nhấn nút Xóa."),
    ("05.6.2", "[CẢI TIẾN v1.1 – A4] Hệ thống kiểm tra: nếu User được chọn chính là Admin đang "
               "đăng nhập, hiển thị lỗi \"Bạn không thể xóa tài khoản của chính mình!\" và dừng."),
    ("05.6.3", "Hệ thống hiển thị hộp thoại xác nhận: \"Xoá user [tên]? Hành động này không thể hoàn tác.\""),
    ("05.6.4", "Admin xác nhận OK → Hệ thống xóa User khỏi CSDL (CASCADE xóa game_sessions liên quan)."),
    ("05.6.5", "Hệ thống load lại danh sách; lùi 1 trang nếu trang hiện tại đã trống."),
    ("05.6.6", "Hệ thống ghi nhận Log vào CSDL với vai trò và trạng thái gốc (action: DELETE_USER)."),
    (None, "Luồng thay thế"),
    ("05.6-A1", "Admin nhấn Huỷ → Đóng dialog, không có thay đổi."),
])

# ══════════════════════════════════════════════════════════════
# UC05.7 – Xem danh sách kết quả
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.7: Xem danh sách kết quả",
               badge="CẢI TIẾN v1.2 – D3",
rows=[
    (None, "Luồng chính"),
    ("05.7.1", "Admin nhấn chọn mục \"Quản lý kết quả\" từ thanh điều hướng."),
    ("05.7.2", "Hệ thống khởi tạo các bộ lọc mặc định: Độ khó = Tất cả, Kết quả = Tất cả, Username trống."),
    ("05.7.3", "Hệ thống truy vấn CSDL và lấy danh sách kết quả, hiển thị tối đa 20 bản ghi / trang."),
    ("05.7.4", "Hệ thống hiển thị danh sách lên bảng kèm: số trang hiện tại, tổng số trang, tổng số kết quả."),
    ("05.7.5", "[CẢI TIẾN v1.2 – D3] Admin có thể nhấn vào tiêu đề cột Score hoặc Thời gian để "
               "sắp xếp tăng/giảm dần (client-side SortedList). Biểu tượng ▲/▼ hiển thị chiều sắp xếp."),
    ("05.7.6", "Admin chọn một trong các chức năng: Lọc kết quả, Xóa kết quả gian lận, "
               "Phát hiện gian lận tự động."),
])

# ══════════════════════════════════════════════════════════════
# UC05.8 – Lọc kết quả
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.8: Lọc kết quả",
rows=[
    (None, "Luồng chính"),
    ("05.8.1", "Admin nhập tên người chơi và/hoặc chọn bộ lọc Độ khó, Kết quả rồi nhấn nút Lọc."),
    ("05.8.2", "Hệ thống xác định điều kiện lọc dựa trên thông tin Admin vừa nhập."),
    ("05.8.3", "Hệ thống truy vấn CSDL theo điều kiện lọc, bắt đầu hiển thị lại từ trang đầu tiên."),
    ("05.8.4", "Hệ thống tải lại bảng với danh sách kết quả phù hợp và cập nhật nhãn số lượng."),
    (None, "Luồng thay thế"),
    ("05.8-A1", "Admin xóa hết điều kiện → Hệ thống nhận diện không có bộ lọc nào, tải lại toàn "
                "bộ danh sách từ CSDL."),
    ("05.8-A2", "Admin nhấn Làm mới → Hệ thống xoá ô Username, reset ComboBox về \"Tất cả\", "
                "tải lại toàn bộ danh sách từ CSDL."),
])

# ══════════════════════════════════════════════════════════════
# UC05.9 – Xóa kết quả gian lận
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.9: Xóa kết quả gian lận",
               badge="CẢI TIẾN v1.2 – B3",
rows=[
    (None, "Luồng chính"),
    ("05.9.1", "Admin tích checkbox trên từng dòng hoặc nhấn \"Chọn tất cả\" để chọn các kết quả "
               "đang hiển thị trên bảng."),
    ("05.9.2", "Admin nhấn nút \"Xoá kết quả gian lận\"."),
    ("05.9.3", "Hệ thống lấy danh sách kết quả mà Admin đã chọn."),
    ("05.9.4", "[CẢI TIẾN v1.2 – B3] Hệ thống hiển thị dialog xác nhận: \"Bạn có chắc muốn xóa "
               "N kết quả đã chọn? Hành động này không thể hoàn tác.\" Admin xác nhận OK."),
    ("05.9.5", "Hệ thống thực hiện xóa toàn bộ các kết quả đã chọn khỏi CSDL."),
    ("05.9.6", "Hệ thống tải lại bảng và hiển thị thông báo thành công kèm số lượng bản ghi đã xóa."),
    ("05.9.7", "Hệ thống ghi nhận log: Admin thực hiện, danh sách mã kết quả bị xóa, tổng số bản ghi "
               "(action: DELETE_SESSION)."),
    (None, "Luồng thay thế"),
    ("05.9-A1", "Admin nhấn Huỷ tại dialog xác nhận → Đóng dialog, không thực hiện xóa."),
])

# ══════════════════════════════════════════════════════════════
# UC05.10 – Phát hiện gian lận tự động [MỚI]
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.10: Phát hiện gian lận tự động",
               badge="MỚI v1.2 – B5",
rows=[
    (None, "Luồng chính"),
    ("05.10.1", "Admin nhấn nút \"Phát hiện gian lận\" trên màn hình Quản lý kết quả."),
    ("05.10.2", "Hệ thống đọc ngưỡng thời gian hợp lý từ bảng game_levels "
                "(ví dụ: EASY < 5s, MEDIUM < 20s, HARD < 60s, EXPERT < 120s)."),
    ("05.10.3", "Hệ thống quét tất cả kết quả WIN đang hiển thị trong bảng, so sánh completion_time "
                "với ngưỡng tương ứng của từng level."),
    ("05.10.4", "Các dòng nghi vấn được tự động highlight màu đỏ nhạt và tích chọn checkbox."),
    ("05.10.5", "Nhãn trạng thái hiển thị: \"Phát hiện N kết quả nghi vấn\"."),
    ("05.10.6", "Admin xem xét và có thể bỏ chọn bất kỳ dòng nào trước khi nhấn Xóa kết quả gian lận "
                "(tiếp tục theo UC05.9)."),
    (None, "Luồng thay thế"),
    ("05.10-A1", "Không tìm thấy kết quả nghi vấn → Hệ thống hiển thị thông báo "
                 "\"Không phát hiện kết quả bất thường nào.\""),
])

# ══════════════════════════════════════════════════════════════
# UC05.11 – Xem thống kê theo người chơi [MỚI]
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.11: Xem thống kê tổng hợp theo người chơi",
               badge="MỚI v1.2 – E1",
rows=[
    (None, "Luồng chính"),
    ("05.11.1", "Admin chọn 1 user trong bảng Quản lý người chơi và nhấn \"Xem chi tiết\" hoặc "
                "double-click vào dòng."),
    ("05.11.2", "Hệ thống truy vấn CSDL: JOIN game_sessions với game_levels theo user_id."),
    ("05.11.3", "Hệ thống hiển thị popup thống kê gồm:\n"
                "  • Thông tin cơ bản: Username, Nickname, Email, Vai trò, Ngày tạo, Lần đăng nhập cuối.\n"
                "  • Thống kê tổng quan: Tổng ván chơi | Tỉ lệ thắng (%) | Tổng điểm tích luỹ.\n"
                "  • Thành tích tốt nhất theo từng độ khó (best_time, best_score).\n"
                "  • Biểu đồ cột điểm số qua các phiên chơi gần nhất (JavaFX BarChart)."),
    ("05.11.4", "Admin xem thông tin và nhấn Đóng để quay lại màn hình chính."),
    (None, "Luồng thay thế"),
    ("05.11-A1", "User chưa có ván chơi nào → Popup hiển thị thông tin cơ bản và thông báo "
                 "\"Người chơi chưa có lịch sử game.\""),
])

# ══════════════════════════════════════════════════════════════
# UC05.12 – Gửi thông báo [MỚI]
# ══════════════════════════════════════════════════════════════
add_flow_table(doc, "Alternative Flow — UC05.12: Gửi thông báo / cảnh báo đến người dùng",
               badge="MỚI v1.2 – E2",
rows=[
    (None, "Luồng chính"),
    ("05.12.1", "Admin chọn 1 hoặc nhiều user trong bảng Quản lý người chơi."),
    ("05.12.2", "Admin nhấn nút \"Gửi thông báo\"."),
    ("05.12.3", "Hệ thống hiển thị dialog soạn thông báo gồm: Tiêu đề (*) và Nội dung (*)."),
    ("05.12.4", "Admin điền nội dung và nhấn Gửi."),
    ("05.12.5", "Hệ thống chèn các bản ghi vào bảng notifications(user_id, title, message, is_read=0, created_at) "
                "cho từng user được chọn."),
    ("05.12.6", "Hệ thống hiển thị thông báo thành công: \"Đã gửi thông báo đến N người dùng\"."),
    ("05.12.7", "Hệ thống ghi nhận Log vào CSDL (action: SEND_NOTIFICATION)."),
    (None, "Luồng thay thế"),
    ("05.12-A1", "Admin nhấn Huỷ → Đóng dialog, không gửi thông báo."),
    ("05.12-A2", "Tiêu đề hoặc Nội dung bị trống → Hệ thống vô hiệu hoá nút Gửi và hiển thị "
                 "gợi ý điền đầy đủ thông tin."),
])

# ══════════════════════════════════════════════════════════════
# EXCEPTION FLOWS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "Exception Flows", 2)

exc_rows = [
    ("05-E1",   "CSDL không thể kết nối → Hệ thống hiển thị hộp thoại lỗi \"Không thể tải dữ liệu\" "
                "và để trống bảng."),
    ("05.3-E1", "Username trùng hoặc lỗi CSDL khi thêm → Hệ thống hiển thị thông báo lỗi cụ thể."),
    ("05.4-E1", "Chưa chọn User, nhấn Chỉnh sửa → Hệ thống hiển thị thông báo \"Hãy chọn user cần sửa\"."),
    ("05.5-E1", "Chưa chọn User, nhấn Khoá → Hệ thống hiển thị thông báo \"Hãy chọn user\"."),
    ("05.5-E2", "[CẢI TIẾN v1.1 – A4] Admin tự khoá mình → Hệ thống hiển thị lỗi "
                "\"Bạn không thể khoá tài khoản của chính mình!\"."),
    ("05.6-E1", "Chưa chọn User, nhấn Xoá → Hệ thống hiển thị thông báo \"Hãy chọn user cần xoá\"."),
    ("05.6-E2", "[CẢI TIẾN v1.1 – A4] Admin tự xoá mình → Hệ thống hiển thị lỗi "
                "\"Bạn không thể xóa tài khoản của chính mình!\"."),
    ("05.9-E1", "Chưa chọn dòng nào, nhấn Xoá → Hiển thị thông báo \"Hãy chọn dữ liệu để xoá\". "
                "Không thực hiện xoá."),
    ("05.12-E1","[MỚI v1.2 – E2] CSDL lỗi khi lưu thông báo → Hiển thị lỗi và không ghi log."),
]

t_exc = make_table(doc, len(exc_rows)+1, 2, col_widths=[3.0, 14.0])
styled_cell(t_exc.cell(0,0), "Mã lỗi",  bold=True, bg=HEADER_BG, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER)
styled_cell(t_exc.cell(0,1), "Mô tả",   bold=True, bg=HEADER_BG, color=(255,255,255))
for i, (code, desc) in enumerate(exc_rows):
    bg = ROW1_BG if i % 2 == 0 else ROW2_BG
    styled_cell(t_exc.cell(i+1,0), code, bold=True, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_cell(t_exc.cell(i+1,1), desc, bg=bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# REVISION HISTORY
# ══════════════════════════════════════════════════════════════
add_heading(doc, "Revision History — Lịch sử thay đổi", 2)

rev_headers = ["Version", "Ngày", "Mục cải tiến", "Mô tả thay đổi"]
rev_data = [
    ("v1.0", "09/05/2026",
     "—",
     "Đặc tả gốc UC05: Basic flow UC05.1; Alternative flows UC05.2 – UC05.9 "
     "(tìm kiếm cơ bản, thêm/sửa/khoá/xoá user, xem/lọc/xoá kết quả gian lận). "
     "Exception flows 05-E1 đến 05.9-E1."),
    ("v1.1", "10/05/2026",
     "Advanced Search; A3; A4",
     "• Advanced Search: Mở rộng UC05.2 bổ sung bộ lọc Email, Ngày tạo (from-to), "
     "Lần đăng nhập gần nhất (from-to), Số ván min/max, Trạng thái có/không có kết quả game.\n"
     "• A3 – Xác nhận khoá/mở khoá: Cập nhật UC05.5, bổ sung bước 05.5.3 hiển thị "
     "dialog xác nhận trước khi thực hiện. Thêm luồng 05.5-A1.\n"
     "• A4 – Chặn Admin tự khoá/xoá: Cập nhật UC05.5 (bước 05.5.2) và UC05.6 (bước 05.6.2); "
     "bổ sung exception 05.5-E2 và 05.6-E2."),
    ("v1.2", "04/06/2026",
     "B3; B5; D1; D2; D3; E1; E2",
     "• B3 – Xác nhận xoá hàng loạt kết quả: Cập nhật UC05.9, bổ sung bước 05.9.4 "
     "dialog xác nhận và luồng 05.9-A1.\n"
     "• B5 – Phát hiện gian lận tự động: Thêm mới UC05.10 (alternative flow hoàn toàn mới).\n"
     "• D1 – Cảnh báo chưa lưu: Cập nhật UC05.3 (05.3-A1) và UC05.4 (05.4-A1) thêm "
     "dialog cảnh báo khi đóng dialog có dữ liệu chưa lưu.\n"
     "• D2 – Badge màu trạng thái: Cập nhật UC05.1 bước 05.1.4, cột Trạng thái hiển thị "
     "🟢/🔴 thay vì text thuần.\n"
     "• D3 – Sắp xếp cột kết quả: Cập nhật UC05.7 bổ sung bước 05.7.5 cho phép sort "
     "theo Score / Thời gian.\n"
     "• E1 – Thống kê theo người chơi: Thêm mới UC05.11 (alternative flow hoàn toàn mới).\n"
     "• E2 – Gửi thông báo: Thêm mới UC05.12 (alternative flow hoàn toàn mới), "
     "bổ sung exception 05.12-E1."),
]

rev_col_widths = [1.5, 2.2, 3.0, 10.3]
t_rev = make_table(doc, len(rev_data)+1, 4, col_widths=rev_col_widths)
for j, h in enumerate(rev_headers):
    styled_cell(t_rev.cell(0,j), h, bold=True, bg=HEADER_BG, color=(255,255,255),
                align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (ver, date, items, desc) in enumerate(rev_data):
    bg = ROW1_BG if i % 2 == 0 else ROW2_BG
    styled_cell(t_rev.cell(i+1,0), ver,   bold=True, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_cell(t_rev.cell(i+1,1), date,  bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_cell(t_rev.cell(i+1,2), items, bold=True, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_cell(t_rev.cell(i+1,3), desc,  bg=bg)

doc.add_paragraph()
p = doc.add_paragraph("Bảng 11: Lịch sử thay đổi đặc tả Use Case UC05")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]; set_font(run, italic=True, size=10)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
out_path = r"D:\BT2_CNPM\Minesweeper\UC05_Dac_Ta_UseCase_v1.2.docx"
doc.save(out_path)
print(f"[OK] Saved: {out_path}")
